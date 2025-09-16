#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
BaiTap1_python: Full script to reproduce the Stata assignment in Python
(Tasks 1..20)

- Usage (from VS Code terminal):
    python baitap1.py --table1 Table1_1.dta --table8 Table8_1.dta --outdir output

- Output: a folder (default "output") with subfolders:
    output/figures    -> png figures
    output/tables     -> csv/html tables
    output/logs       -> program log + textual model summaries
    output/reports    -> a simple Word report (report.docx) and pickled models

Dependencies (install first):
    pip install pandas numpy matplotlib seaborn statsmodels scipy python-docx tabulate

Author: Generated for UEH assignment (Gujarati 2014 datasets)
"""

import os
import sys
import argparse
import logging
from datetime import datetime
import textwrap

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns

import statsmodels.api as sm
import statsmodels.formula.api as smf
from statsmodels.stats.outliers_influence import variance_inflation_factor
from statsmodels.stats.diagnostic import het_breuschpagan
from statsmodels.stats.outliers_influence import OLSInfluence
from scipy import stats

from docx import Document
from docx.shared import Inches
from tabulate import tabulate

# ----------------------------- helpers ---------------------------------

def ensure_dir(path):
    if not os.path.exists(path):
        os.makedirs(path, exist_ok=True)


def setup_logging(logfile_path):
    logger = logging.getLogger()
    logger.setLevel(logging.INFO)
    # console
    ch = logging.StreamHandler(sys.stdout)
    ch.setLevel(logging.INFO)
    fmt = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
    ch.setFormatter(fmt)
    logger.addHandler(ch)
    # file
    fh = logging.FileHandler(logfile_path, mode='w', encoding='utf-8')
    fh.setLevel(logging.INFO)
    fh.setFormatter(fmt)
    logger.addHandler(fh)
    return logger


def save_df_outputs(df, name, out_tables_dir):
    csv_path = os.path.join(out_tables_dir, f"{name}.csv")
    html_path = os.path.join(out_tables_dir, f"{name}.html")
    df.to_csv(csv_path, index=True)
    try:
        df.to_html(html_path, index=True, float_format="{:.6g}".format)
    except Exception:
        # fallback
        with open(html_path, 'w', encoding='utf-8') as f:
            f.write(df.to_csv(index=True))
    return csv_path, html_path


def save_text(text, path):
    with open(path, 'w', encoding='utf-8') as f:
        f.write(text)


# ----------------------------- plotting ---------------------------------

def plot_histogram(series, title, outpath, bins=30, show_kde=True):
    plt.figure(figsize=(8,5))
    sns.histplot(series.dropna(), bins=bins, kde=show_kde)
    plt.title(title)
    plt.xlabel(series.name)
    plt.tight_layout()
    plt.savefig(outpath, dpi=300)
    plt.close()


def plot_scatter_with_fit(df, x, y, outpath, annotate_corr=True):
    plt.figure(figsize=(8,6))
    sns.regplot(x=x, y=y, data=df, scatter_kws={'s':20}, line_kws={'linewidth':2})
    plt.xlabel(x)
    plt.ylabel(y)
    plt.title(f"Scatter: {y} vs {x} with linear fit")
    if annotate_corr:
        corr = df[[x,y]].dropna().corr().iloc[0,1]
        plt.annotate(f'corr={corr:.3f}', xy=(0.05,0.95), xycoords='axes fraction')
    plt.tight_layout()
    plt.savefig(outpath, dpi=300)
    plt.close()


# ----------------------------- econometrics -----------------------------

def ols_regression(df, formula, robust=False):
    model = smf.ols(formula, data=df).fit()
    if robust:
        model_robust = smf.ols(formula, data=df).fit(cov_type='HC3')
        return model, model_robust
    return model, None


def compute_vif(df, features):
    X = df[features].copy()
    X = X.assign(const=1)
    cols = X.columns.tolist()
    vif_data = []
    for i in range(len(cols)-1):  # omit last const for VIF calculation duplicates
        col = cols[i]
        try:
            vif = variance_inflation_factor(X.values, i)
        except Exception:
            vif = np.nan
        vif_data.append((col, vif))
    return pd.DataFrame(vif_data, columns=['feature','VIF']).set_index('feature')


def breusch_pagan_test(model):
    exog = model.model.exog
    lm, lm_pvalue, fvalue, f_pvalue = het_breuschpagan(model.resid, exog)
    return {'lm': lm, 'lm_pvalue': lm_pvalue, 'fvalue': fvalue, 'f_pvalue': f_pvalue}


# ----------------------------- logistic/probit margins -----------------

def logistic_pred_prob(params, x):
    # logistic function
    z = np.dot(x, params)
    return 1.0 / (1.0 + np.exp(-z))


def logit_marginal_effects_at_point(params, point):
    # params: pandas Series (including intercept), point: dict or pandas Series with variable values
    # model uses linear index: intercept + beta_j * x_j
    # For interactions, supply corresponding values in point (i.e. ageedu = age*edu etc.)
    beta = params
    # ensure same ordering: intercept first if named 'Intercept' or 'const'
    if 'Intercept' in beta.index:
        const_name = 'Intercept'
    elif 'const' in beta.index:
        const_name = 'const'
    else:
        const_name = beta.index[0]
    x = []
    for name in beta.index:
        if name == const_name:
            x.append(1.0)
        else:
            x.append(point.get(name, 0.0))
    x = np.array(x)
    z = np.dot(x, beta.values)
    p = 1.0/(1.0+np.exp(-z))
    # For each variable j (not intercept): marginal = (d z / d x_j) * p*(1-p)
    # where d z / d x_j = beta_j plus contributions from interaction terms that include x_j.
    # We'll compute derivatives numerically by constructing a small delta change for each variable.
    me = {}
    eps = 1e-5
    for name in beta.index:
        if name == const_name:
            continue
        x_plus = x.copy()
        idx = list(beta.index).index(name)
        x_plus[idx] += eps
        z_plus = np.dot(x_plus, beta.values)
        p_plus = 1.0/(1.0+np.exp(-z_plus))
        derivative = (p_plus - p)/eps
        me[name] = derivative
    return {'p': p, 'marginal_effects': me}


def probit_marginal_effects_at_point(params, point):
    # For probit: marginal = (d z / d x_j) * phi(z)
    if 'Intercept' in params.index:
        const_name = 'Intercept'
    elif 'const' in params.index:
        const_name = 'const'
    else:
        const_name = params.index[0]
    x = []
    for name in params.index:
        if name == const_name:
            x.append(1.0)
        else:
            x.append(point.get(name, 0.0))
    x = np.array(x)
    z = np.dot(x, params.values)
    phi = 1/np.sqrt(2*np.pi)*np.exp(-0.5*z*z)
    me = {}
    eps = 1e-5
    for name in params.index:
        if name == const_name:
            continue
        idx = list(params.index).index(name)
        x_plus = x.copy()
        x_plus[idx] += eps
        z_plus = np.dot(x_plus, params.values)
        derivative_z = (z_plus - z)/eps
        me[name] = derivative_z * phi
    return {'z': z, 'phi': phi, 'marginal_effects': me}


# ----------------------------- main workflow ---------------------------

def main(args):
    # prepare folders
    OUT = args.outdir
    ensure_dir(OUT)
    FIG_DIR = os.path.join(OUT, 'figures')
    TABLE_DIR = os.path.join(OUT, 'tables')
    LOG_DIR = os.path.join(OUT, 'logs')
    REPORT_DIR = os.path.join(OUT, 'reports')
    MODEL_DIR = os.path.join(OUT, 'models')
    for d in [FIG_DIR, TABLE_DIR, LOG_DIR, REPORT_DIR, MODEL_DIR]:
        ensure_dir(d)

    logfile = os.path.join(LOG_DIR, f'baitap1_python_{datetime.now().strftime("%Y%m%d_%H%M%S")}.log')
    logger = setup_logging(logfile)
    logger.info('Starting bai tap 1 pipeline')

    # ------------------ Part 1: Table1_1.dta ---------------------------
    logger.info('Loading Table1_1 data from %s', args.table1)
    df1 = pd.read_stata(args.table1)
    logger.info('Data loaded: %d rows, %d columns', df1.shape[0], df1.shape[1])
    save_df_outputs(df1.head(100), 'table1_head', TABLE_DIR)

    # 1. describe data
    logger.info('1) Describe the dataset (info + dtypes + head)')
    descr_text = textwrap.dedent( (
        f"Data shape: {df1.shape}\n\n" +
        f"Dtypes:\n{df1.dtypes}\n\n" +
        f"Head:\n{df1.head().to_string()}\n"
    ) )
    save_text(descr_text, os.path.join(LOG_DIR, 'part1_describe.txt'))

    # 2. summary stats
    logger.info('2) Summary statistics')
    summ = df1.describe(include='all')
    save_df_outputs(summ, 'part1_summary', TABLE_DIR)

    # 3. mean wage by gender
    logger.info('3) Mean wage by female')
    mean_by_female = df1.groupby('female')['wage'].agg(['count','mean','std','min','max']).rename(columns={'mean':'mean_wage'})
    save_df_outputs(mean_by_female, 'mean_wage_by_female', TABLE_DIR)

    # 4. mean wage by female & nonwhite
    logger.info('4) Mean wage by female and nonwhite')
    mean_by_two = df1.groupby(['female','nonwhite'])['wage'].agg(['count','mean','std']).rename(columns={'mean':'mean_wage'})
    save_df_outputs(mean_by_two, 'mean_wage_by_female_nonwhite', TABLE_DIR)

    # 5. cross-tab female x nonwhite with percentages
    logger.info('5) Crosstab female x nonwhite (counts + row% + col% + total%)')
    ct = pd.crosstab(df1['female'], df1['nonwhite'])
    ct_row = pd.crosstab(df1['female'], df1['nonwhite'], normalize='index')*100
    ct_col = pd.crosstab(df1['female'], df1['nonwhite'], normalize='columns')*100
    save_df_outputs(ct, 'crosstab_counts', TABLE_DIR)
    save_df_outputs(ct_row, 'crosstab_rowpct', TABLE_DIR)
    save_df_outputs(ct_col, 'crosstab_colpct', TABLE_DIR)

    # 6. histograms for wage and lnwage
    logger.info('6) Histogram wage and lnwage')
    # make sure wage > 0 for log
    df1['lnwage'] = np.where(df1['wage']>0, np.log(df1['wage']), np.nan)
    plot_histogram(df1['wage'], 'Histogram of wage', os.path.join(FIG_DIR, 'hist_wage.png'))
    plot_histogram(df1['lnwage'], 'Histogram of ln(wage)', os.path.join(FIG_DIR, 'hist_lnwage.png'))

    # 7. scatter education -> wage
    logger.info('7) Scatter wage vs education')
    plot_scatter_with_fit(df1, x='education', y='wage', outpath=os.path.join(FIG_DIR, 'scatter_wage_education.png'))
    plot_scatter_with_fit(df1, x='education', y='lnwage', outpath=os.path.join(FIG_DIR, 'scatter_lnwage_education.png'))

    # 8. t-test wage difference between male/female
    logger.info('8) t-test wage by female')
    groups = df1.dropna(subset=['wage','female']).groupby('female')['wage']
    try:
        male = groups.get_group(0)
        female = groups.get_group(1)
        tstat, pval = stats.ttest_ind(male, female, equal_var=False, nan_policy='omit')
        ttest_text = f"t-statistic={tstat:.6g}, p-value={pval:.6g}, n_male={len(male)}, n_female={len(female)}"
    except Exception as e:
        ttest_text = f"t-test failed: {e}"
    save_text(ttest_text, os.path.join(LOG_DIR, 'ttest_female_wage.txt'))

    # 9. ANOVA across age groups (we construct age = exper + education + 6)
    logger.info('9) ANOVA across age groups')
    # create age
    df1['age'] = df1.get('exper', 0) + df1.get('education', 0) + 6
    df1['agegroup'] = pd.cut(df1['age'], bins=[0,30,40,50,60,100], labels=['<30','30-39','40-49','50-59','60+'])
    anova_groups = [g['wage'].dropna().values for name,g in df1.groupby('agegroup')]
    # remove empty groups
    anova_groups = [g for g in anova_groups if len(g)>0]
    try:
        F, p = stats.f_oneway(*anova_groups)
        anova_text = f"ANOVA F-statistic={F:.6g}, p-value={p:.6g}"
    except Exception as e:
        anova_text = f"ANOVA failed: {e}"
    save_text(anova_text, os.path.join(LOG_DIR, 'anova_agegroups.txt'))

    # 10. OLS regression
    logger.info('10) OLS regression: wage ~ female + nonwhite + union + education + exper')
    formula = 'wage ~ female + nonwhite + union + education + exper'
    model, model_robust = ols_regression(df1, formula, robust=True)
    model_summary_text = model.summary().as_text()
    save_text(model_summary_text, os.path.join(LOG_DIR, 'ols_summary.txt'))
    if model_robust is not None:
        save_text(model_robust.summary().as_text(), os.path.join(LOG_DIR, 'ols_summary_robust.txt'))
    # save params
    pd.DataFrame({'coef': model.params, 'std_err': model.bse, 't': model.tvalues, 'p': model.pvalues}).to_csv(os.path.join(TABLE_DIR, 'ols_params.csv'))

    # 11. predictions
    logger.info('11) Predict fitted values and residuals')
    df1['yhat'] = model.predict(df1)
    df1['resid'] = model.resid
    df1[['wage','yhat','resid']].to_csv(os.path.join(TABLE_DIR, 'ols_pred_resid.csv'), index=False)

    # scatter actual vs predicted
    plt.figure(figsize=(7,6))
    sns.scatterplot(x='yhat', y='wage', data=df1, s=20)
    plt.plot([df1['yhat'].min(), df1['yhat'].max()], [df1['yhat'].min(), df1['yhat'].max()], color='red', linestyle='--')
    plt.xlabel('Predicted wage')
    plt.ylabel('Actual wage')
    plt.title('Actual vs Predicted wage')
    plt.tight_layout()
    plt.savefig(os.path.join(FIG_DIR, 'actual_vs_predicted_wage.png'), dpi=300)
    plt.close()

    # 12. joint test: all regressors = 0
    logger.info('12) Joint F-test for all regressors excluded')
    try:
        f_test = model.f_test('female = nonwhite = union = education = exper = 0')
        ftest_text = str(f_test)
    except Exception as e:
        ftest_text = f'F-test failed: {e}'
    save_text(ftest_text, os.path.join(LOG_DIR, 'ols_joint_test.txt'))

    # 13. multicollinearity and heteroskedasticity
    logger.info('13) VIF and Breusch-Pagan')
    features = ['female','nonwhite','union','education','exper']
    vif_df = compute_vif(df1.dropna(subset=features), features)
    save_df_outputs(vif_df, 'vif', TABLE_DIR)

    bp = breusch_pagan_test(model)
    save_text(str(bp), os.path.join(LOG_DIR, 'breusch_pagan.txt'))

    # influence measures (Cook's D)
    influence = OLSInfluence(model)
    cooks, pvals = influence.cooks_distance
    df1['cooks_d'] = cooks
    df1.sort_values('cooks_d', ascending=False)[['wage','yhat','cooks_d']].head(20).to_csv(os.path.join(TABLE_DIR, 'top20_cooksd.csv'))

    # ------------------ Part 2: Table8_1.dta (logistic & probit) ---------
    logger.info('Loading Table8_1 data from %s', args.table8)
    df2 = pd.read_stata(args.table8)
    logger.info('Data2 loaded: %d rows, %d columns', df2.shape[0], df2.shape[1])
    save_df_outputs(df2.head(100), 'table8_head', TABLE_DIR)

        # 14. frequency of smoker
    logger.info('14) Frequency table for smoker')
    smoker_freq = df2['smoker'].value_counts(dropna=False).to_frame('count')
    smoker_freq['pct'] = df2['smoker'].value_counts(normalize=True, dropna=False).values * 100
    save_df_outputs(smoker_freq, 'smoker_frequency', TABLE_DIR)

    # 15. logistic regression
    # chuẩn hóa tên cột để dùng 'edu' cho tiện
    if 'educ' in df2.columns and 'edu' not in df2.columns:
        df2 = df2.rename(columns={'educ': 'edu'})

    # tạo biến tương tác nếu chưa có trong dataset
    if 'ageedu' not in df2.columns:
        df2['ageedu'] = df2['age'] * df2['edu']
    if 'educincome' not in df2.columns:
        df2['educincome'] = df2['edu'] * df2['income']

    logger.info('15) Logit: smoker ~ age + edu + income + pcigs79 + ageedu + educincome')
    logit_formula = 'smoker ~ age + edu + income + pcigs79 + ageedu + educincome'
    logit_model = smf.logit(logit_formula, data=df2).fit(disp=False)
    save_text(logit_model.summary().as_text(), os.path.join(LOG_DIR, 'logit_summary.txt'))
    pd.DataFrame({
        'coef': logit_model.params,
        'se': logit_model.bse,
        'z': logit_model.tvalues,
        'p': logit_model.pvalues
    }).to_csv(os.path.join(TABLE_DIR, 'logit_params.csv'))

    # 16. predict probability for each respondent
    logger.info('16) Predict probability for each respondent')
    df2['phat_logit'] = logit_model.predict(df2)
    df2[['smoker','phat_logit']].head(50).to_csv(os.path.join(TABLE_DIR, 'logit_pred_prob_sample.csv'))

    # 17. marginal effects at mean
    logger.info('17) Marginal effects at means (logit)')
    try:
        mfx_mean = logit_model.get_margeff(at='mean', method='dydx')
        save_text(mfx_mean.summary().as_text(), os.path.join(LOG_DIR, 'logit_mfx_atmean.txt'))
        # also save as table
        mfx_df = mfx_mean.summary_frame()
        save_df_outputs(mfx_df, 'logit_mfx_atmean', TABLE_DIR)
    except Exception as e:
        save_text(f"mfx mean failed: {e}", os.path.join(LOG_DIR, 'logit_mfx_atmean.txt'))

    # 18. marginal effects at specific point (age=63, edu=10, income=20000, pcigs79=60)
    logger.info('18) Marginal effects at specific point (age=63, edu=10, income=20000, pcigs79=60)')
    point = {
        'age': 63.0,
        'edu': 10.0,
        'income': 20000.0,
        'pcigs79': 60.0,
        # interaction terms
        'ageedu': 63.0 * 10.0,
        'educincome': 10.0 * 20000.0
    }
    logit_me_point = logit_marginal_effects_at_point(logit_model.params, point)
    save_text(str(logit_me_point), os.path.join(LOG_DIR, 'logit_mfx_point.txt'))

        # -------------------------
    # 19. Probit model
    # -------------------------
    logger.info('19) Probit model')
    # ensure logit_formula exists (defined earlier)
    try:
        probit_model = smf.probit(logit_formula, data=df2).fit(disp=False)
        save_text(probit_model.summary().as_text(), os.path.join(LOG_DIR, 'probit_summary.txt'))
        pd.DataFrame({
            'coef': probit_model.params,
            'se': probit_model.bse,
            'z': probit_model.tvalues,
            'p': probit_model.pvalues
        }).to_csv(os.path.join(TABLE_DIR, 'probit_params.csv'))
    except Exception as e:
        save_text(f'Probit estimation failed: {e}', os.path.join(LOG_DIR, 'probit_summary.txt'))
        logger.exception('Probit estimation failed')
        probit_model = None

    # marginal effects for probit at means
    if probit_model is not None:
        try:
            probit_mfx_mean = probit_model.get_margeff(at='mean', method='dydx')
            save_text(probit_mfx_mean.summary().as_text(), os.path.join(LOG_DIR, 'probit_mfx_atmean.txt'))
            save_df_outputs(probit_mfx_mean.summary_frame(), 'probit_mfx_atmean', TABLE_DIR)
        except Exception as e:
            save_text(f"probit mfx failed: {e}", os.path.join(LOG_DIR, 'probit_mfx_atmean.txt'))
            logger.exception('Probit marginal effects failed')

        # marginal effects at point for probit (analytical)
        try:
            probit_me_point = probit_marginal_effects_at_point(probit_model.params, point)
            save_text(str(probit_me_point), os.path.join(LOG_DIR, 'probit_mfx_point.txt'))
        except Exception as e:
            save_text(f"probit mfx point failed: {e}", os.path.join(LOG_DIR, 'probit_mfx_point.txt'))
            logger.exception('Probit mfx at point failed')

    # -------------------------
    # 20. Compare logit vs probit
    # -------------------------
    logger.info('20) Compare logit vs probit')
    try:
        # create coefficient comparison table
        if (logit_model is not None) and (probit_model is not None):
            comp = pd.concat([logit_model.params.rename('logit_coef'), probit_model.params.rename('probit_coef')], axis=1)
        elif logit_model is not None:
            comp = pd.DataFrame({'logit_coef': logit_model.params})
        elif probit_model is not None:
            comp = pd.DataFrame({'probit_coef': probit_model.params})
        else:
            comp = pd.DataFrame()

        if not comp.empty:
            save_df_outputs(comp, 'logit_probit_coeffs', TABLE_DIR)

        # compare marginal effects if available
        try:
            if 'mfx_df' in locals() and 'probit_mfx_mean' in locals():
                pmfx_df = probit_mfx_mean.summary_frame()
                merged_mfx = pd.concat([mfx_df['dydx'].rename('logit_dydx'), pmfx_df['dydx'].rename('probit_dydx')], axis=1)
                save_df_outputs(merged_mfx, 'compare_marginal_effects', TABLE_DIR)
        except Exception:
            pass
    except Exception as e:
        logger.exception('Error when comparing logit and probit')
        save_text(f'Compare error: {e}', os.path.join(LOG_DIR, 'compare_error.txt'))

    # ------------------ Create a polished Word report ---------------------
    logger.info('Creating a polished Word report at %s', REPORT_DIR)
    try:
        doc = Document()
        # Title / cover
        doc.add_heading('BÀI TẬP 1 — Phân tích bằng Python', level=0)
        doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph('Dữ liệu: Table1_1.dta (OLS) và Table8_1.dta (Logit/Probit)')
        doc.add_page_break()

        # Part 1: Descriptive & OLS
        doc.add_heading('Phần 1: Thống kê mô tả & Hồi quy OLS', level=1)
        doc.add_paragraph('Kết quả thống kê mô tả và OLS được lưu trong thư mục `output/tables` và biểu đồ trong `output/figures`.')

        # Insert key figures if exist
        if os.path.exists(os.path.join(FIG_DIR, 'hist_wage.png')):
            doc.add_paragraph('Histogram of wage:')
            doc.add_picture(os.path.join(FIG_DIR, 'hist_wage.png'), width=Inches(6))
        if os.path.exists(os.path.join(FIG_DIR, 'scatter_wage_education.png')):
            doc.add_paragraph('Scatter plot: wage vs education')
            doc.add_picture(os.path.join(FIG_DIR, 'scatter_wage_education.png'), width=Inches(6))

        # Insert short OLS summary (first lines)
        doc.add_heading('OLS: wage ~ female + nonwhite + union + education + exper', level=2)
        try:
            ols_txt = model.summary().as_text()  # model exists earlier for OLS
            # put summary in small font paragraph
            p = doc.add_paragraph()
            run = p.add_run(ols_txt)
            run.font.size = Pt(8)
        except Exception:
            doc.add_paragraph('OLS summary saved in logs/ols_summary.txt')

        doc.add_page_break()

        # Part 2: Logit & Probit
        doc.add_heading('Phần 2: Logit & Probit', level=1)
        doc.add_paragraph('Mô hình hồi quy nhị phân và so sánh kết quả.')

        # Logit results table
        logit_path = os.path.join(TABLE_DIR, 'logit_params.csv')
        if os.path.exists(logit_path):
            logit_df = pd.read_csv(logit_path, index_col=0)
            doc.add_heading('Logit regression (coefficients)', level=2)
            t = doc.add_table(rows=1, cols=len(logit_df.columns) + 1)
            t.style = 'Light List Accent 1'
            hdr = t.rows[0].cells
            hdr[0].text = 'Variable'
            for j, col in enumerate(logit_df.columns, start=1):
                hdr[j].text = col
            for var, row in logit_df.iterrows():
                rc = t.add_row().cells
                rc[0].text = str(var)
                for j, col in enumerate(logit_df.columns, start=1):
                    val = row[col]
                    rc[j].text = f'{val:.4f}' if pd.notnull(val) else ''

        # Probit results table
        probit_path = os.path.join(TABLE_DIR, 'probit_params.csv')
        if os.path.exists(probit_path):
            probit_df = pd.read_csv(probit_path, index_col=0)
            doc.add_heading('Probit regression (coefficients)', level=2)
            t = doc.add_table(rows=1, cols=len(probit_df.columns) + 1)
            t.style = 'Light List Accent 2'
            hdr = t.rows[0].cells
            hdr[0].text = 'Variable'
            for j, col in enumerate(probit_df.columns, start=1):
                hdr[j].text = col
            for var, row in probit_df.iterrows():
                rc = t.add_row().cells
                rc[0].text = str(var)
                for j, col in enumerate(probit_df.columns, start=1):
                    val = row[col]
                    rc[j].text = f'{val:.4f}' if pd.notnull(val) else ''

        doc.add_page_break()

        # Odds ratios & short interpretation
        if logit_model is not None:
            doc.add_heading('Odds Ratios (Logit)', level=2)
            ors = np.exp(logit_model.params)
            t = doc.add_table(rows=1, cols=2)
            t.style = 'Medium Shading 1 Accent 1'
            hdr = t.rows[0].cells
            hdr[0].text = 'Variable'
            hdr[1].text = 'Odds Ratio (exp(beta))'
            for var, orval in ors.items():
                rc = t.add_row().cells
                rc[0].text = str(var)
                rc[1].text = f'{orval:.4f}'

            doc.add_paragraph(
                "Interpretation: Odds Ratio > 1 indicates an increase in odds of being a smoker when the variable increases by 1 unit; "
                "Odds Ratio < 1 indicates a decrease."
            )

        # Comparison text
        doc.add_heading('So sánh & Nhận xét ngắn', level=2)
        doc.add_paragraph(
            "Kết quả Logit và Probit có dấu hệ số tương tự. Marginal effects (xem output/tables) cung cấp ảnh hưởng trực tiếp lên xác suất. "
            "Khi có heteroskedasticity hoặc mẫu nhỏ, hãy đối chiếu margin và kiểm định robust."
        )

        # Save report
        report_path = os.path.join(REPORT_DIR, 'report.docx')
        doc.save(report_path)
        logger.info('Report saved to %s', report_path)
    except Exception as e:
        logger.exception('Failed to create Word report')
        save_text(f'Report generation failed: {e}', os.path.join(LOG_DIR, 'report_error.txt'))

    logger.info('All done. Outputs are in the folder: %s', OUT)
    logger.info('Log file saved to %s', logfile)


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Run BaiTap1 analysis in Python')
    parser.add_argument('--table1', type=str, default='Table1_1.dta', help='Path to Table1_1.dta')
    parser.add_argument('--table8', type=str, default='Table8_1.dta', help='Path to Table8_1.dta')
    parser.add_argument('--outdir', type=str, default='output', help='Output directory')
    args = parser.parse_args()
    main(args)
    
